from docx import Document
from typing import List, Dict, Any
import re
from pypdf import PdfReader
import pdfplumber
import io
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from app.util.openai_client import get_image_caption
import fitz
import os


page_out_dir = os.path.join(os.getcwd(), "data", "image_crops")

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    texts = [para.text for para in doc.paragraphs if para.text.strip()]

    for table in doc.tables:
        headers = [cell.text for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            first_col = row.cells[0].text.strip() if len(row.cells) > 0 else ''
            for col_idx, cell in enumerate(row.cells):
                if col_idx == 0:
                    continue
                header = headers[col_idx] if col_idx < len(headers) else f"Column {col_idx+1}"
                value = cell.text.strip()
                if value:
                    cell_text = f"Table: {header} | Row: {first_col}: | Value: {value}"
                    texts.append(cell_text)
    return '\n'.join(texts)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using pypdf.
    Tries to preserve simple paragraph seperation. Ignores images.
    """
    reader = PdfReader(file_path)
    pages_text = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt:
            txt = re.sub(r"\n{2,}", "\n\n", txt.strip())
            pages_text.append(txt)
    return "\n\n".join(pages_text)

def extract_rich_pdf_segments(file_path: str, text_extract_only: bool = False) -> List[Dict[str, Any]]:
    segments: List[Dict[str, Any]]= []    
    fitz_doc = fitz.open(file_path)
    img_idx = 0
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_index, page in enumerate(pdf.pages):                
                page_num = page_index + 1
                try:            
                    text = page.extract_text() or ""                    
                except Exception as ex:                    
                    text = ""
                if text.strip():
                    paras = [p.strip() for p in re.split(r"\n{2,}",text) if p.strip()]
                    for para in paras:
                        segments.append({'type': 'text', 'page': page_num, 'content': para})                                    
                try:                    
                    tables = page.extract_tables() or []
                except Exception as ex:                    
                    tables = []                
                for table in tables:
                    if not table or len(table) < 2:
                        continue
                    header = table[0]
                    for row in table[1:]:
                        row_extended = row + ["" for _ in range(len(header)-len(row))]
                        cells = []
                        for h, v in zip(header, row_extended):
                            h_clean = (h or '').strip()
                            v_clean = (v or '').strip()
                            if v_clean:
                                cells.append(f"{h_clean}:{v_clean}" if h_clean else v_clean)
                        if cells:
                            segments.append({'type': 'table_row', 'page': page_num, 'content': ' | '.join(cells)})                                
                if text_extract_only:
                    continue                
                image_regions = []
                
                #try:
                #    imgs = convert_from_path(file_path, first_page=page_num, last_page=page_num, dpi=200)[0]                                       
                #except Exception:                                        
                #    imgs = []
                #for img in imgs:
                #    x0= img.get('x0'); y0 = img.get('y0'); x1 = img.get('x1'); y1= img.get('y1')
                #    name = img.get('name') or img.get('object_type') or 'image'
                #    image_regions.append({
                #        'x0': x0, 'y0': y0, 'x1': x1, 'y1': y1,
                #        'origin': 'bottom-left', 'name': name
                #    })
                
                
                try:
                    if fitz_doc is not None:
                        fitz_page = fitz_doc[page_index]
                    else:
                        fitx_page = None
                    if fitz_page is not None:
                        raw = fitz_page.get_text("rawdict") or {}
                        for blk in raw.get('blocks', []):
                            if blk.get('type') == 1:                                
                                bbox = blk.get('bbox') or []
                                print(len(bbox))
                                if len(bbox) == 4:
                                    x0, y0, x1, y1 = bbox
                                    image_regions.append({
                                        'x0': x0, 'y0': y0, 'x1': x1, 'y1': y1,
                                        'origin': 'top-left', 'name': 'image'
                                    })
                        if not image_regions:
                            try:
                                img_list = fitz_page.get_images(full=True) or []                                
                                for info in img_list:
                                    xref = info[0]
                                    rects = fitz_page.get_image_rects(xref) or []                                    
                                    for r in rects:
                                        image_regions.append({
                                            'x0': float(r.x0), 'y0': float(r.y0), 'x1': float(r.x1), 'y1': float(r.y1),
                                            'origin': 'top-left', 'name': f'image:{xref}'
                                        })
                            except Exception:
                                pass
                        if not image_regions:
                            image_regions.append({'x0':0.0, 'y0': 0.0, 'x1': float(page.width), 'y1': float(page.height), 'origin': 'top-left', 'name':'page_full'})
                except Exception:
                    pass
                
                if image_regions:
                    try:
                        page_img = convert_from_path(file_path, first_page=page_num, last_page=page_num, dpi=200)[0]
                    except Exception as e:
                        import traceback
                        traceback.print_exc()
                        print(e)
                        page_img = None
                    for region in image_regions:
                        x0 = region['x0']; y0= region['y0']; x1 = region['x1']; y1= region['y1']; 
                        name = region.get('name', 'image')
                        origin = region.get('origin', 'bottom-left')                    
                        if page_img and all(v is not None for v in [x0,y0,x1,y1]):
                            #print("i am inside")
                            try:
                                pg_w, pg_h = float(page.width), float(page.height)
                                
                                fx0 = max(0.0, min(1.0, float(x0)/pg_w))
                                fy0 = max(0.0, min(1.0, float(y0)/pg_h))
                                fx1 = max(0.0, min(1.0, float(x1)/pg_w))
                                fy1 = max(0.0, min(1.0, float(y1)/pg_h))
                                
                                img_w, img_h = page_img.size
                                if origin == 'bottom-left':
                                    top = int((1.0-fy1) * img_h)
                                    bottom = int((1.0-fy0) * img_h)
                                else:
                                    top = int(fy0 * img_h)
                                    bottom = int(fy1 * img_h)
                                left = int(fx0 * img_w)
                                right = int(fx1 * img_w)
                                left, right = max(0, min(left, img_w)), max(0, min(right, img_w))
                                top, bottom = max(0, min(top, img_h)), max(0, min(bottom, img_h))
                                #print(f'right - {right}, left - {left}, bottom -{bottom}, top-{top}')
                                if right > left and bottom > top:                                                                
                                    crop = page_img.crop((left, top, right, bottom))
                                    img_idx += 1
                                    fname = f"img_{page_num:04d}_{img_idx:03d}.png"
                                    out_path = os.path.join(page_out_dir, fname)
                                    buf = io.BytesIO()                                   
                                    #crop.save(out_path, format='PNG')
                                    crop.save(buf, format='PNG')
                                    caption = get_image_caption(buf.getvalue())
                                    print(caption)
                                    content = f"name={name}, caption={caption}"
                                    segments.append({'type':'image', 'page':page_num, 'content': content})
                                else:
                                    meta = f"name={name}, bbox=({x0},{y0},{x1},{y1})"
                                    segments.append({'type':'image', 'page':page_num, 'content': meta})
                            except Exception as e:
                                import traceback
                                traceback.print_exc()
                                meta = f"name={name}, bbox=({x0},{y0},{x1},{y1})"
                                segments.append({'type':'image', 'page':page_num, 'content': meta})
                        else:
                            meta = f"name={name}, bbox=({x0},{y0},{x1},{y1})"
                            segments.append({'type':'image', 'page':page_num, 'content': meta})
    except Exception as e:        
        import traceback
        traceback.print_exc()
        fallback = extract_text_from_pdf(file_path)
        if fallback.strip():
            segments.append({'type': 'text', 'page': 1, 'content': fallback.strip()})
    return segments

def chunk_segments(segments: List[Dict[str, Any]], max_words: int = 300) -> List[str]:
    chunks: List[str] = []
    current_words = 0
    current_parts: List[str] = []
    current_mode: str | None = None

    def flush():
        nonlocal current_parts, current_words, current_mode
        if current_parts:
            chunks.append('\n'.join(current_parts).strip())
        current_parts = []
        current_words = 0
        current_mode = None

    for seg in segments:
        seg_words = len(seg['content'].split())
        mode = seg['type']
        prefix = f"[PAGE {seg['page']} | {mode.upper()}]"
        line = prefix + seg['content']
        if current_mode is not None and mode !=current_mode:
            flush()
        if current_words + seg_words > max_words and current_parts:
            flush()
        current_parts.append(line)
        current_words += seg_words
        current_mode = mode
    flush()
    return chunks

def chunk_text(text: str, max_chunk_size: int = 1000) -> List[str]:
    sentences = re.split(r'(?<=[.!?]) +', text)
    chunks = []
    current = ""
    for sent in sentences:
        if len((current+sent).split()) > max_chunk_size:
            chunks.append(current.strip())
            current = sent + " "
        else:
            current += sent + " "
    if current.split():
        chunks.append(current.strip())
    return chunks
